#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для подготовки данных из .docx документа для векторизации и хранения в ChromaDB
с использованием LangChain и улучшенного TextSplitter
"""

import os
import uuid
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from pathlib import Path

# Импорты для обработки документов
from docx import Document

# Импорты LangChain
from langchain.schema import Document as LangChainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.embeddings.base import Embeddings

# Альтернативные эмбеддинги (если нет OpenAI)
from langchain_community.embeddings import HuggingFaceEmbeddings

# Импорт нашего улучшенного сплиттера
from improved_text_splitter import ImprovedTextSplitter

@dataclass
class DocumentChunk:
    """Класс для представления фрагмента документа"""
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    
class DocxToChromaDB:
    """Класс для обработки .docx файлов и сохранения в ChromaDB"""
    
    def __init__(self, 
                 embedding_model: str = "intfloat/multilingual-e5-large",
                 chromadb_path: str = "./chromadb",
                 collection_name: str = "documents"):
        """
        Инициализация процессора
        
        Args:
            embedding_model: Модель для создания эмбеддингов
            chromadb_path: Путь к базе данных ChromaDB
            collection_name: Имя коллекции в ChromaDB
        """
        self.embedding_model = embedding_model
        self.chromadb_path = chromadb_path
        self.collection_name = collection_name
        self.improved_splitter = ImprovedTextSplitter()
        
        # Создаем директорию для ChromaDB если её нет
        os.makedirs(chromadb_path, exist_ok=True)
        
        # Инициализируем эмбеддинги
        self.embeddings = self._init_embeddings()
        
    def _init_embeddings(self) -> Embeddings:
        """Инициализация модели эмбеддингов"""
        try:
            # Пытаемся использовать OpenAI embeddings если есть API ключ
            if os.getenv("OPENAI_API_KEY"):
                print("Использую OpenAI embeddings")
                return OpenAIEmbeddings()
            else:
                print(f"Использую HuggingFace embeddings: {self.embedding_model}")
                return HuggingFaceEmbeddings(model_name=self.embedding_model)
        except Exception as e:
            print(f"Ошибка при инициализации эмбеддингов: {e}")
            print(f"Использую HuggingFace embeddings: {self.embedding_model}")
            return HuggingFaceEmbeddings(model_name=self.embedding_model)
    
    def read_docx(self, file_path: str) -> str:
        """Читает содержимое .docx файла"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Обрабатываем таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            print(f"Ошибка при чтении .docx файла: {e}")
            return ""
    
    def create_chunks(self, text: str, source_file: str) -> List[DocumentChunk]:
        """Создает фрагменты документа с метаданными"""
        chunks = []
        
        # Используем улучшенный сплиттер для структурной разбивки
        parts = self.improved_splitter.split_text(text, add_headers=True)
        
        for i, part in enumerate(parts):
            chunk_id = str(uuid.uuid4())
            
            # Создаем метаданные
            metadata = {
                "source": source_file,
                "chunk_id": chunk_id,
                "chunk_index": i,
                "type": part.get("type", "content"),
                "total_chunks": len(parts)
            }
            
            # Добавляем структурную информацию
            if "header" in part:
                metadata["header"] = part["header"]
            
            # Добавляем размер фрагмента
            metadata["chunk_size"] = len(part["content"])
            
            chunks.append(DocumentChunk(
                content=part["content"],
                metadata=metadata,
                chunk_id=chunk_id
            ))
        
        return chunks
    
    def create_additional_chunks(self, text: str, source_file: str, 
                               chunk_size: int = 1000, 
                               chunk_overlap: int = 200) -> List[DocumentChunk]:
        """Создает дополнительные фрагменты с помощью RecursiveCharacterTextSplitter"""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = []
        texts = text_splitter.split_text(text)
        
        for i, chunk_text in enumerate(texts):
            chunk_id = str(uuid.uuid4())
            
            metadata = {
                "source": source_file,
                "chunk_id": chunk_id,
                "chunk_index": i,
                "type": "recursive_split",
                "total_chunks": len(texts),
                "chunk_size": len(chunk_text),
                "splitter_type": "recursive"
            }
            
            chunks.append(DocumentChunk(
                content=chunk_text,
                metadata=metadata,
                chunk_id=chunk_id
            ))
        
        return chunks
    
    def save_to_chromadb(self, chunks: List[DocumentChunk], 
                        collection_name: Optional[str] = None) -> Chroma:
        """Сохраняет фрагменты в ChromaDB"""
        if collection_name is None:
            collection_name = self.collection_name
            
        # Подготавливаем документы для LangChain
        documents = []
        for chunk in chunks:
            doc = LangChainDocument(
                page_content=chunk.content,
                metadata=chunk.metadata
            )
            documents.append(doc)
        
        # Создаем или обновляем векторную базу данных
        try:
            vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                collection_name=collection_name,
                persist_directory=self.chromadb_path
            )
            
            # В ChromaDB 0.4.x+ данные автоматически сохраняются
            # vectorstore.persist() больше не нужен
            
            print(f"Сохранено {len(documents)} фрагментов в ChromaDB")
            print(f"Коллекция: {collection_name}")
            print(f"Путь к БД: {self.chromadb_path}")
            
            return vectorstore
            
        except Exception as e:
            print(f"Ошибка при сохранении в ChromaDB: {e}")
            return None
    
    def process_docx_file(self, file_path: str, 
                         use_both_splitters: bool = True,
                         recursive_chunk_size: int = 1000,
                         recursive_chunk_overlap: int = 200) -> Dict[str, Any]:
        """
        Полная обработка .docx файла
        
        Args:
            file_path: Путь к .docx файлу
            use_both_splitters: Использовать оба сплиттера (структурный и рекурсивный)
            recursive_chunk_size: Размер фрагмента для рекурсивного сплиттера
            recursive_chunk_overlap: Перекрытие для рекурсивного сплиттера
            
        Returns:
            Словарь с результатами обработки
        """
        print(f"Обрабатываем файл: {file_path}")
        
        # Читаем файл
        text = self.read_docx(file_path)
        if not text:
            print("Не удалось прочитать файл!")
            return {"error": "Не удалось прочитать файл"}
        
        print(f"Извлечено {len(text)} символов")
        
        # Создаем фрагменты с помощью улучшенного сплиттера
        structured_chunks = self.create_chunks(text, file_path)
        print(f"Создано {len(structured_chunks)} структурных фрагментов")
        
        # Сохраняем структурные фрагменты
        vectorstore_structured = self.save_to_chromadb(
            structured_chunks, 
            f"{self.collection_name}_structured"
        )
        
        result = {
            "file_path": file_path,
            "text_length": len(text),
            "structured_chunks": len(structured_chunks),
            "vectorstore_structured": vectorstore_structured
        }
        
        # Создаем дополнительные фрагменты с рекурсивным сплиттером
        if use_both_splitters:
            recursive_chunks = self.create_additional_chunks(
                text, file_path, recursive_chunk_size, recursive_chunk_overlap
            )
            print(f"Создано {len(recursive_chunks)} рекурсивных фрагментов")
            
            # Сохраняем рекурсивные фрагменты
            vectorstore_recursive = self.save_to_chromadb(
                recursive_chunks, 
                f"{self.collection_name}_recursive"
            )
            
            result.update({
                "recursive_chunks": len(recursive_chunks),
                "vectorstore_recursive": vectorstore_recursive
            })
        
        return result
    
    def search_similar(self, query: str, 
                      collection_name: Optional[str] = None,
                      k: int = 5) -> List[Dict[str, Any]]:
        """Поиск похожих фрагментов"""
        if collection_name is None:
            collection_name = f"{self.collection_name}_structured"
        
        try:
            # Загружаем векторную базу данных
            vectorstore = Chroma(
                collection_name=collection_name,
                embedding_function=self.embeddings,
                persist_directory=self.chromadb_path
            )
            
            # Выполняем поиск
            results = vectorstore.similarity_search_with_score(query, k=k)
            
            formatted_results = []
            for doc, score in results:
                formatted_results.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "similarity_score": score
                })
            
            return formatted_results
            
        except Exception as e:
            print(f"Ошибка при поиске: {e}")
            return []

def main():
    """Основная функция для демонстрации работы"""
    # Инициализируем процессор
    processor = DocxToChromaDB(
        embedding_model="intfloat/multilingual-e5-large",
        chromadb_path="./chromadb_documents",
        collection_name="contract_documents"
    )
    
    # Обрабатываем client.docx
    docx_file = "client.docx"
    
    if not os.path.exists(docx_file):
        print(f"Файл {docx_file} не найден!")
        return
    
    print("=== ОБРАБОТКА ДОКУМЕНТА ===")
    result = processor.process_docx_file(
        docx_file, 
        use_both_splitters=True,
        recursive_chunk_size=800,
        recursive_chunk_overlap=100
    )
    
    print("\n=== РЕЗУЛЬТАТЫ ===")
    for key, value in result.items():
        if not key.startswith("vectorstore"):
            print(f"{key}: {value}")
    
    print("\n=== ТЕСТОВЫЙ ПОИСК ===")
    
    # Примеры поиска
    test_queries = [
        "лизинговые платежи",
        "страхование объекта",
        "права лизингодателя",
        "договор поставки"
    ]
    
    for query in test_queries:
        print(f"\nПоиск по запросу: '{query}'")
        results = processor.search_similar(query, k=3)
        
        for i, result in enumerate(results, 1):
            print(f"\n{i}. Похожесть: {result['similarity_score']:.4f}")
            print(f"Тип: {result['metadata'].get('type', 'N/A')}")
            if 'header' in result['metadata']:
                print(f"Заголовок: {result['metadata']['header']}")
            print(f"Содержимое: {result['content'][:200]}...")
    
    print("\n=== СТАТИСТИКА КОЛЛЕКЦИЙ ===")
    
    # Проверяем количество документов в коллекциях
    try:
        import chromadb
        client = chromadb.PersistentClient(path="./chromadb_documents")
        
        structured_collection = client.get_collection("contract_documents_structured")
        recursive_collection = client.get_collection("contract_documents_recursive")
        
        print(f"Структурных фрагментов: {structured_collection.count()}")
        print(f"Рекурсивных фрагментов: {recursive_collection.count()}")
        
    except Exception as e:
        print(f"Ошибка при получении статистики: {e}")

if __name__ == "__main__":
    main() 
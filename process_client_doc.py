#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для обработки файла client.doc с помощью text_splitter.py
"""

import os
import sys
from text_splitter import TextSplitter

def read_doc_file(file_path):
    """Читает содержимое .doc файла"""
    try:
        # Пытаемся использовать python-docx
        try:
            from docx import Document
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except ImportError:
            print("Модуль python-docx не установлен, пытаемся установить...")
            os.system("pip install python-docx")
            from docx import Document
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
    except Exception as e:
        print(f"Ошибка при чтении .doc файла: {e}")
        print("Попробуем прочитать как текстовый файл...")
        
        # Попытка чтения как текстового файла
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except:
            try:
                with open(file_path, 'r', encoding='cp1251') as f:
                    return f.read()
            except:
                try:
                    with open(file_path, 'r', encoding='latin1') as f:
                        return f.read()
                except Exception as final_e:
                    print(f"Не удалось прочитать файл: {final_e}")
                    return None

def process_client_doc():
    """Основная функция обработки client.doc"""
    print("=== ОБРАБОТКА ФАЙЛА client.doc ===\n")
    
    # Путь к файлу
    doc_path = "client.doc"
    
    # Проверяем существование файла
    if not os.path.exists(doc_path):
        print(f"Файл {doc_path} не найден!")
        return
    
    # Читаем содержимое файла
    print("Читаем содержимое файла...")
    content = read_doc_file(doc_path)
    
    if content is None:
        print("Не удалось прочитать файл!")
        return
    
    print(f"Файл успешно прочитан. Размер: {len(content)} символов\n")
    
    # Выводим первые 500 символов для проверки
    print("=== ПЕРВЫЕ 500 СИМВОЛОВ ФАЙЛА ===")
    print(content[:500])
    print("...\n")
    
    # Создаем экземпляр text_splitter
    splitter = TextSplitter()
    
    # Обрабатываем текст
    print("Обрабатываем текст с помощью TextSplitter...")
    parts = splitter.split_text(content, add_headers=True)
    
    print(f"Текст разбит на {len(parts)} частей\n")
    
    # Выводим результат
    print("=== РЕЗУЛЬТАТ ОБРАБОТКИ ===")
    formatted_output = splitter.format_output(parts)
    print(formatted_output)
    
    # Сохраняем результат в файл
    output_file = "client_doc_processed.txt"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("=== ОБРАБОТАННЫЙ ФАЙЛ client.doc ===\n\n")
        f.write(f"Исходный файл: {doc_path}\n")
        f.write(f"Дата обработки: {os.popen('date').read().strip()}\n")
        f.write(f"Размер исходного текста: {len(content)} символов\n")
        f.write(f"Количество частей: {len(parts)}\n\n")
        f.write("=== ОБРАБОТАННЫЙ ТЕКСТ ===\n\n")
        f.write(formatted_output)
    
    print(f"\nРезультат сохранен в файл: {output_file}")
    
    # Показываем структурный анализ
    print("\n" + "="*60)
    print("СТРУКТУРНЫЙ АНАЛИЗ:")
    print("="*60)
    
    structured_output = splitter.format_output(parts, show_structure=True)
    print(structured_output)
    
    # Сохраняем структурный анализ
    structure_file = "client_doc_structure.txt"
    with open(structure_file, 'w', encoding='utf-8') as f:
        f.write("=== СТРУКТУРНЫЙ АНАЛИЗ client.doc ===\n\n")
        f.write(structured_output)
    
    print(f"\nСтруктурный анализ сохранен в файл: {structure_file}")
    
    # Статистика
    print("\n=== СТАТИСТИКА ===")
    sections = [p for p in parts if p['type'] == 'section']
    subsections = [p for p in parts if p['type'] == 'subsection']
    subpoints = [p for p in parts if p['type'] == 'subpoint']
    content_parts = [p for p in parts if p['type'] == 'content']
    
    print(f"Разделов: {len(sections)}")
    print(f"Подразделов: {len(subsections)}")
    print(f"Подпунктов: {len(subpoints)}")
    print(f"Содержательных частей: {len(content_parts)}")
    
    # Показываем заголовки разделов
    if sections:
        print("\n=== НАЙДЕННЫЕ РАЗДЕЛЫ ===")
        for i, section in enumerate(sections, 1):
            print(f"{i}. {section['content']}")
    
    if subsections:
        print("\n=== НАЙДЕННЫЕ ПОДРАЗДЕЛЫ ===")
        for i, subsection in enumerate(subsections, 1):
            print(f"{i}. {subsection['content']}")

def main():
    """Главная функция"""
    try:
        process_client_doc()
    except Exception as e:
        print(f"Произошла ошибка: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 
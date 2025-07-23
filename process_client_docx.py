#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для обработки файла client.docx с помощью text_splitter.py
"""

import os
import sys
import re
from text_splitter import TextSplitter

def read_docx_file(file_path):
    """Читает содержимое .docx файла"""
    try:
        # Устанавливаем python-docx если не установлен
        try:
            from docx import Document
        except ImportError:
            print("Устанавливаем python-docx...")
            os.system("pip install python-docx")
            from docx import Document
        
        print(f"Открываем файл {file_path}...")
        doc = Document(file_path)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        # Обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return '\n'.join(text_parts) if text_parts else ""
    
    except Exception as e:
        print(f"Ошибка при чтении .docx файла: {e}")
        return None

def clean_extracted_text(text):
    """Очищает извлеченный текст"""
    if not text:
        return ""
    
    # Удаляем лишние пробелы
    text = re.sub(r'\s+', ' ', text)
    
    # Удаляем строки короче 3 символов
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if len(line) > 3:
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def analyze_document_structure(parts):
    """Анализирует структуру документа"""
    sections = [p for p in parts if p['type'] == 'section']
    subsections = [p for p in parts if p['type'] == 'subsection']
    subpoints = [p for p in parts if p['type'] == 'subpoint']
    content_parts = [p for p in parts if p['type'] == 'content']
    
    return {
        'sections': sections,
        'subsections': subsections,
        'subpoints': subpoints,
        'content_parts': content_parts
    }

def save_results(original_text, processed_parts, analysis):
    """Сохраняет результаты обработки"""
    # Основной результат
    splitter = TextSplitter()
    formatted_output = splitter.format_output(processed_parts)
    
    # Сохраняем основной результат
    with open("client_docx_processed.txt", "w", encoding="utf-8") as f:
        f.write("=== ОБРАБОТАННЫЙ ФАЙЛ client.docx ===\n\n")
        f.write(f"Дата обработки: {os.popen('date').read().strip()}\n")
        f.write(f"Размер исходного текста: {len(original_text)} символов\n")
        f.write(f"Количество частей: {len(processed_parts)}\n\n")
        f.write("=== ИСХОДНЫЙ ТЕКСТ ===\n\n")
        f.write(original_text)
        f.write("\n\n=== ОБРАБОТАННЫЙ ТЕКСТ ===\n\n")
        f.write(formatted_output)
    
    # Сохраняем структурный анализ
    structured_output = splitter.format_output(processed_parts, show_structure=True)
    with open("client_docx_structure.txt", "w", encoding="utf-8") as f:
        f.write("=== СТРУКТУРНЫЙ АНАЛИЗ client.docx ===\n\n")
        f.write(structured_output)
    
    # Сохраняем только чистый текст
    clean_content = []
    for part in processed_parts:
        if part['type'] == 'content':
            if 'header' in part:
                clean_content.append(f"[{part['header']}]")
            clean_content.append(part['content'])
        elif part['type'] in ['section', 'subsection', 'subpoint']:
            clean_content.append(f"\n{part['content']}")
    
    with open("client_docx_clean.txt", "w", encoding="utf-8") as f:
        f.write("=== ЧИСТЫЙ ТЕКСТ client.docx ===\n\n")
        f.write('\n'.join(clean_content))
    
    return formatted_output, structured_output

def main():
    """Основная функция"""
    print("=== ОБРАБОТКА ФАЙЛА client.docx ===\n")
    
    file_path = "client.docx"
    
    # Проверяем существование файла
    if not os.path.exists(file_path):
        print(f"Файл {file_path} не найден!")
        return
    
    # Читаем содержимое файла
    print("Извлекаем текст из файла...")
    content = read_docx_file(file_path)
    
    if not content:
        print("Не удалось извлечь текст из файла!")
        return
    
    print(f"Извлечено {len(content)} символов")
    
    # Очищаем текст
    cleaned_content = clean_extracted_text(content)
    print(f"После очистки: {len(cleaned_content)} символов")
    
    # Показываем первые 500 символов
    print("\n=== ПЕРВЫЕ 500 СИМВОЛОВ ===")
    print(cleaned_content[:500])
    if len(cleaned_content) > 500:
        print("...")
    
    # Создаем экземпляр TextSplitter
    splitter = TextSplitter()
    
    # Обрабатываем текст
    print("\n=== ОБРАБОТКА С ПОМОЩЬЮ TextSplitter ===")
    parts = splitter.split_text(cleaned_content, add_headers=True)
    
    print(f"Текст разбит на {len(parts)} частей")
    
    # Анализируем структуру
    structure_analysis = analyze_document_structure(parts)
    
    # Сохраняем результаты
    formatted_output, structured_output = save_results(cleaned_content, parts, structure_analysis)
    
    print("\n=== РЕЗУЛЬТАТ ОБРАБОТКИ ===")
    print(formatted_output[:1000])
    if len(formatted_output) > 1000:
        print("...")
    
    # Выводим статистику
    print("\n=== СТАТИСТИКА ===")
    print(f"Разделов: {len(structure_analysis['sections'])}")
    print(f"Подразделов: {len(structure_analysis['subsections'])}")
    print(f"Подпунктов: {len(structure_analysis['subpoints'])}")
    print(f"Содержательных частей: {len(structure_analysis['content_parts'])}")
    
    # Показываем найденные разделы
    if structure_analysis['sections']:
        print("\n=== НАЙДЕННЫЕ РАЗДЕЛЫ ===")
        for i, section in enumerate(structure_analysis['sections'], 1):
            print(f"{i}. {section['content']}")
    
    if structure_analysis['subsections']:
        print("\n=== НАЙДЕННЫЕ ПОДРАЗДЕЛЫ ===")
        for i, subsection in enumerate(structure_analysis['subsections'], 1):
            print(f"{i}. {subsection['content']}")
    
    if structure_analysis['subpoints']:
        print("\n=== НАЙДЕННЫЕ ПОДПУНКТЫ ===")
        for i, subpoint in enumerate(structure_analysis['subpoints'], 1):
            print(f"{i}. {subpoint['content']}")
    
    print("\n=== СОХРАНЕННЫЕ ФАЙЛЫ ===")
    print("1. client_docx_processed.txt - полный результат обработки")
    print("2. client_docx_structure.txt - структурный анализ")
    print("3. client_docx_clean.txt - чистый текст")
    
    print("\n=== ОБРАБОТКА ЗАВЕРШЕНА ===")

if __name__ == "__main__":
    main() 